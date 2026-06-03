import docx, sys, os
from docx.shared import Pt
sys.stdout.reconfigure(encoding='utf-8')

src = r'D:\Desktop\ctccbz\Web технологии\WEB-технологии\ИТз-211 Поляков Курсовая Web-технологии.docx'
dst = r'D:\Desktop\ctccbz\Web технологии\WEB-технологии\ИТз-211 Поляков Курсовая Web-технологии_fixed_v2.docx'

doc = docx.Document(src)

# Step 1: Find the appendix code-listing boundary dynamically.
# The "Models" section heading marks where the foreign/regenerated listings begin.
cut = None
for i, p in enumerate(doc.paragraphs):
    if i > 600 and p.text.strip() == 'Models':
        cut = i
        break
if cut is None:
    raise SystemExit('Boundary heading "Models" not found - aborting to avoid damage.')

print(f'Boundary "Models" found at paragraph {cut}')
print(f'  prev: |{doc.paragraphs[cut-1].text[:60]}|  (should be blank)')
print(f'  prev2: |{doc.paragraphs[cut-2].text[:60]}|  (end of InteractiveExplanationController)')

to_remove = [doc.paragraphs[i]._element for i in range(cut, len(doc.paragraphs))]
print(f'Removing {len(to_remove)} paragraphs from boundary to end')
for elem in to_remove:
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)

base = r'D:\Docs\Codegpt\kurs'

def add_section_heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = 'Times New Roman'

def add_file(rel_path, display_name):
    filepath = os.path.join(base, rel_path)
    if not os.path.exists(filepath):
        print(f'  SKIP (not found): {display_name}')
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        code = f.read()
    # filename sub-heading
    p = doc.add_paragraph()
    r = p.add_run(display_name)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'
    # code lines
    for line in code.split('\n'):
        pc = doc.add_paragraph()
        rc = pc.add_run(line)
        rc.font.size = Pt(10)
        rc.font.name = 'Courier New'
    print(f'  Added: {display_name} ({len(code)} chars)')

# Step 2: Models (Eloquent) — Article.php is already shown above, skip it
add_section_heading('Модели Eloquent (app/Models/)')
models = [
    ('app/Models/User.php', 'User.php'),
    ('app/Models/Category.php', 'Category.php'),
    ('app/Models/Tag.php', 'Tag.php'),
    ('app/Models/Comment.php', 'Comment.php'),
    ('app/Models/Rating.php', 'Rating.php'),
    ('app/Models/DailyUpdate.php', 'DailyUpdate.php'),
    ('app/Models/InteractiveExplanation.php', 'InteractiveExplanation.php'),
]
for rel, name in models:
    add_file(rel, name)

# Step 3: Controllers (key ones) — InteractiveExplanationController already above, skip it
add_section_heading('Контроллеры (app/Http/Controllers/)')
controllers = [
    ('app/Http/Controllers/HomeController.php', 'HomeController.php'),
    ('app/Http/Controllers/ArticleController.php', 'ArticleController.php'),
    ('app/Http/Controllers/SearchController.php', 'SearchController.php'),
    ('app/Http/Controllers/CommentController.php', 'CommentController.php'),
    ('app/Http/Controllers/RatingController.php', 'RatingController.php'),
    ('app/Http/Controllers/FavoriteController.php', 'FavoriteController.php'),
    ('app/Http/Controllers/DailyUpdateController.php', 'DailyUpdateController.php'),
    ('app/Http/Controllers/Admin/ArticleController.php', 'Admin/ArticleController.php'),
    ('app/Http/Controllers/Admin/WikipediaImportController.php', 'Admin/WikipediaImportController.php'),
]
for rel, name in controllers:
    add_file(rel, name)

# Step 3b: Routes — web.php already shown at top of appendix, so add auth.php
add_section_heading('Маршруты (routes/)')
add_file('routes/auth.php', 'auth.php')

# Step 4: Views (Vue components) — unique ones only
add_section_heading('Представления View (resources/js/)')
views = [
    ('resources/js/Layouts/AppLayout.vue', 'Layouts/AppLayout.vue'),
    ('resources/js/Components/InteractiveMode.vue', 'Components/InteractiveMode.vue'),
    ('resources/js/Components/DailyUpdateBlock.vue', 'Components/DailyUpdateBlock.vue'),
    ('resources/js/Components/TipTapEditor.vue', 'Components/TipTapEditor.vue'),
]
for rel, name in views:
    add_file(rel, name)

try:
    doc.save(dst)
    print(f'\nSaved to: {dst}')
except PermissionError:
    alt = dst.replace('_fixed_v2.docx', '_fixed_v3.docx')
    doc.save(alt)
    print(f'\n(v2 was locked) Saved to: {alt}')
